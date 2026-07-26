#!/usr/bin/env python3
"""Render the integration specs from Markdown to formatted Word documents.

    python scripts/md2docx.py            # rebuild every document in MANIFEST
    python scripts/md2docx.py overview   # rebuild one, by key

Requires `python-docx` (pip install python-docx). No other dependency — pandoc
and LibreOffice are deliberately not required, since neither is installed on the
build machines.

Supported Markdown: headings, paragraphs, tables (including \\| escapes), fenced
code blocks, blockquotes, bullet / numbered / checkbox lists, horizontal rules,
and inline bold, italic, code and links.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

# ── what gets built ──────────────────────────────────────────────────────────
MANIFEST = {
    # The master dossier — every integration document in one file.
    # Deliberately EXCLUDES: INTEGRATION-TAX-SYSTEM.md and
    # INTEGRATION-MDA-REVENUE-APP.md (already contained in the combined spec, so
    # including them would duplicate ~2,500 lines), plus DEPLOYMENT-RUNBOOK.md
    # and BRIS-BUILD-STATUS.md (internal operations, not interface material).
    'master': dict(
        srcs=[
            ('docs/INTEGRATION-OVERVIEW.md',
             'Overview — how the two revenue systems fit together'),
            ('docs/INTEGRATION-REVENUE-DATA-REQUIREMENTS.md',
             'Requirements — the data interface specification'),
            ('docs/INTEGRATION-PAYE.md',
             'Annex A — PAYE sync, existing push contract'),
            ('docs/INTEGRATION-TAX-PAYMENTS.md',
             'Annex B — tax payments sync, existing push contract'),
            ('docs/COMPLIANCE-MAPPING.md',
             'Annex C — compliance control mapping'),
            ('docs/LEGAL-BRIEF-PROVIDER-PENALTIES.md',
             'Annex D — legal brief on provider penalties'),
        ],
        title='BIZDATA — Revenue Integration Dossier',
        frontmatter=[
            ('Document type', 'Complete integration dossier'),
            ('Contents', 'Overview · requirements specification · four annexes'),
            ('Counterparties', 'Tax Administration System and MDA Revenue Application'),
            ('Direction', '**Pull** — BIZDATA calls each system on a schedule'),
            ('Scope', 'State-administered taxes and non-tax revenue'),
            ('Status', 'Draft for review'),
            ('Version', '1.0'),
        ],
        out='BIZDATA-Revenue-Integration-Dossier.docx',
        subtitle='Revenue Integration Dossier',
        tagline='Every integration document in one file — overview, requirements '
                'specification, and supporting annexes',
        meta='Version 1.0   ·   Draft for review\n'
             'Overview · Requirements · Annexes A–D\n'
             'State-administered revenue only',
    ),
    'combined': dict(
        src='docs/INTEGRATION-REVENUE-DATA-REQUIREMENTS.md',
        out='BIZDATA-Revenue-Data-Requirements.docx',
        subtitle='Revenue Data Interface Requirements',
        tagline='What BIZDATA needs from the authority’s revenue systems, and why',
        meta='Version 1.0   ·   Draft for review\n'
             'Part C — Tax Administration System   ·   Part D — MDA Revenue Application\n'
             'Part A — management   ·   Parts B–G — engineering',
    ),
    'tax-system': dict(
        src='docs/INTEGRATION-TAX-SYSTEM.md',
        out='BIZDATA-A-Tax-System-Data-Requirements.docx',
        subtitle='Document A — Tax Administration System',
        tagline='Data interface requirements: what FINDATA needs the Tax System '
                'to expose, and why',
        meta='Version 2.0   ·   Draft for review by the Tax System team\n'
             'Integration model: BIZDATA pulls on a schedule\n'
             'Part A — management   ·   Parts B–E — engineering',
    ),
    'mda-revenue-app': dict(
        src='docs/INTEGRATION-MDA-REVENUE-APP.md',
        out='BIZDATA-B-MDA-Revenue-App-Data-Requirements.docx',
        subtitle='Document B — MDA Revenue Application',
        tagline='Data interface requirements: what BIZDATA needs the MDA Revenue '
                'Application to expose, and why',
        meta='Version 1.0   ·   Draft for review by the MDA Revenue Application team\n'
             'Integration model: BIZDATA pulls on a schedule\n'
             'Part A — management   ·   Parts B–E — engineering',
    ),
    'overview': dict(
        src='docs/INTEGRATION-OVERVIEW.md',
        out='BIZDATA-Revenue-Integration-Overview.docx',
        subtitle='Revenue Integration Overview',
        tagline='How the two counterparty documents fit together',
        meta='Version 1.0   ·   Internal / joint overview\n'
             'Companions: Document A (Tax System) and Document B (MDA Revenue Application)',
    ),
}

# ── house style ──────────────────────────────────────────────────────────────
MONO = 'Consolas'
BODY = 'Calibri'
ACCENT = RGBColor(0x0F, 0x62, 0x5C)      # brand teal
MUTED = RGBColor(0x77, 0x88, 0x88)
CODE_FILL = 'F4F6F6'
HDR_FILL = '0F625C'
ALT_FILL = 'F7F9F9'


# ── low-level docx helpers ───────────────────────────────────────────────────
def shade_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def left_bar(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pbdr.append(left)
    pPr.append(pbdr)


def bottom_rule(p, color='D6DEDE'):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pbdr.append(bot)
    pPr.append(pbdr)


def keep_with_next(p):
    p._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))


def add_field(paragraph, instr):
    """Insert a Word field code (used for the page number)."""
    r = paragraph.add_run()
    beg = OxmlElement('w:fldChar'); beg.set(qn('w:fldCharType'), 'begin')
    r._r.append(beg)
    r2 = paragraph.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r3._r.append(end)


# ── inline markdown ──────────────────────────────────────────────────────────
INLINE = re.compile(
    r'(\*\*.+?\*\*)'           # bold
    r'|(`[^`]+`)'              # code
    r'|(\[[^\]]+\]\([^)]+\))'  # link
    r'|(\*[^*\n]+\*)'          # italic
)


def _run(p, txt, size=None, color=None, bold=False, italic=False, mono=False, link=False):
    if not txt:
        return
    r = p.add_run(txt)
    r.font.name = MONO if mono else BODY
    if size:
        r.font.size = size
    if mono:
        r.font.size = Pt((size.pt if size else 10.5) - 1)
        r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
    elif color:
        r.font.color.rgb = color
    if link:
        r.font.color.rgb = ACCENT
        r.font.underline = True
    r.bold = bold
    r.italic = italic


def add_inline(p, text, base_size=None, base_color=None):
    text = text.replace('\\|', '|')
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(p, text[pos:m.start()], base_size, base_color)
        tok = m.group(0)
        if tok.startswith('**'):
            _run(p, tok[2:-2], base_size, base_color, bold=True)
        elif tok.startswith('`'):
            _run(p, tok[1:-1], base_size, base_color, mono=True)
        elif tok.startswith('['):
            label = tok[1:tok.index(']')]
            target = tok[tok.index('(') + 1:-1]
            # external links stay links; internal doc refs render as italic text
            _run(p, label, base_size, base_color,
                 link=target.startswith('http'), italic=not target.startswith('http'))
        else:
            _run(p, tok[1:-1], base_size, base_color, italic=True)
        pos = m.end()
    if pos < len(text):
        _run(p, text[pos:], base_size, base_color)


# ── block builders ───────────────────────────────────────────────────────────
def add_code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.18)
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 0)
        pf.line_spacing = 1.0
        shade_para(p, CODE_FILL)
        r = p.add_run(line if line else ' ')
        r.font.name = MONO
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x2A)


def split_row(line):
    cells = re.split(r'(?<!\\)\|', line.strip())
    if cells and not cells[0].strip():
        cells = cells[1:]
    if cells and not cells[-1].strip():
        cells = cells[:-1]
    return [c.strip() for c in cells]


def add_table(doc, rows):
    header, body = rows[0], rows[2:]
    ncols = len(header)
    t = doc.add_table(rows=1, cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, h, base_size=Pt(9))
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, HDR_FILL)
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for i in range(ncols):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, row[i] if i < len(row) else '', base_size=Pt(9))
            if ri % 2 == 1:
                shade_cell(cells[i], ALT_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_cover(doc, subtitle, tagline, meta_lines):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BIZDATA')
    r.font.size = Pt(38); r.bold = True; r.font.color.rgb = ACCENT; r.font.name = BODY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(17); r.font.color.rgb = RGBColor(0x33, 0x44, 0x44); r.font.name = BODY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tagline)
    r.font.size = Pt(11.5); r.italic = True; r.font.color.rgb = MUTED; r.font.name = BODY

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta_lines)
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55, 0x66, 0x66); r.font.name = BODY
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ── source loading ───────────────────────────────────────────────────────────
def load_markdown(spec):
    """Return one Markdown string for a spec, whether single-source or a dossier.

    For a multi-source spec each document's own `# Title` line is replaced by the
    label given in the manifest, so it renders as a titled divider rather than
    competing with the dossier's own title.
    """
    if 'src' in spec:
        return (ROOT / spec['src']).read_text(encoding='utf-8')

    parts = [f"# {spec['title']}\n"]
    if spec.get('frontmatter'):
        parts.append('\n| | |\n|---|---|\n')
        parts += [f'| {k} | {v} |\n' for k, v in spec['frontmatter']]
    for rel, label in spec['srcs']:
        text = (ROOT / rel).read_text(encoding='utf-8')
        text = re.sub(r'\A\s*#\s+.*$', f'# {label}', text, count=1, flags=re.M)
        parts.append('\n\n' + text.rstrip() + '\n')
    return ''.join(parts)


# ── main conversion ──────────────────────────────────────────────────────────
def convert(md_text, out_path, subtitle='', tagline='', meta_lines=''):
    src = md_text.split('\n')

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.15

    for lvl, size in ((1, 20), (2, 15), (3, 12.5), (4, 11)):
        s = doc.styles[f'Heading {lvl}']
        s.font.name = BODY
        s.font.size = Pt(size)
        s.font.color.rgb = ACCENT
        s.font.bold = True
        s.paragraph_format.space_before = Pt(16 if lvl <= 2 else 11)
        s.paragraph_format.space_after = Pt(5)

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.85)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(f'BIZDATA  ·  {subtitle}  ·  Page ')
    add_field(fp, ' PAGE ')
    for r in fp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
        r.font.name = BODY

    build_cover(doc, subtitle, tagline, meta_lines)

    i, n, seen_title = 0, len(src), False
    while i < n:
        line = src[i]

        if line.startswith('```'):
            i += 1
            buf = []
            while i < n and not src[i].startswith('```'):
                buf.append(src[i]); i += 1
            i += 1
            add_code_block(doc, buf)
            continue

        if line.startswith('|') and i + 1 < n and re.match(r'^\|[\s:|-]+\|?$', src[i + 1].strip()):
            rows = []
            while i < n and src[i].startswith('|'):
                rows.append(split_row(src[i])); i += 1
            if len(rows) >= 2:
                add_table(doc, rows)
            continue

        stripped = line.strip()

        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            bottom_rule(p)
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            lvl, text = len(m.group(1)), m.group(2)
            if lvl == 1 and not seen_title:      # the doc title is on the cover
                seen_title = True
                i += 1
                continue
            if lvl == 1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            h = doc.add_heading(level=min(lvl, 4))
            h.text = ''
            add_inline(h, text)
            for r in h.runs:
                r.font.color.rgb = ACCENT
                r.bold = True
                r.font.name = BODY
            keep_with_next(h)
            i += 1
            continue

        if stripped.startswith('>'):
            buf = []
            while i < n and src[i].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', src[i])); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(7)
            left_bar(p, '0F625C')
            add_inline(p, ' '.join(x.strip() for x in buf if x.strip()),
                       base_size=Pt(9.5), base_color=RGBColor(0x3A, 0x4A, 0x4A))
            continue

        m = re.match(r'^-\s+\[([ xX])\]\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(3)
            p.add_run('☐  ').font.size = Pt(11)
            body = m.group(2)
            j = i + 1
            while j < n and src[j].startswith('      ') and src[j].strip():
                body += ' ' + src[j].strip(); j += 1
            add_inline(p, body)
            i = j
            continue

        m = re.match(r'^([-*])\s+(.*)$', stripped)
        if m:
            indent = len(line) - len(line.lstrip())
            body = m.group(2)
            j = i + 1
            while (j < n and src[j].strip()
                   and not re.match(r'^\s*([-*]|\d+\.)\s', src[j])
                   and not src[j].startswith(('|', '```'))
                   and (len(src[j]) - len(src[j].lstrip())) > indent):
                body += ' ' + src[j].strip(); j += 1
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3 + 0.25 * (indent // 2))
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, body)
            i = j
            continue

        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            indent = len(line) - len(line.lstrip())
            body = m.group(2)
            j = i + 1
            while (j < n and src[j].strip()
                   and not re.match(r'^\s*(\d+\.|[-*])\s', src[j])
                   and not src[j].startswith(('|', '```'))
                   and (len(src[j]) - len(src[j].lstrip())) > indent):
                body += ' ' + src[j].strip(); j += 1
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, body)
            i = j
            continue

        if not stripped:
            i += 1
            continue

        buf = [stripped]
        j = i + 1
        while j < n and src[j].strip() and not re.match(
                r'^\s*(#{1,6}\s|[-*]\s|\d+\.\s|>|\||```|---$)', src[j]):
            buf.append(src[j].strip()); j += 1
        add_inline(doc.add_paragraph(), ' '.join(buf))
        i = j

    doc.save(out_path)
    return out_path


def build(keys=None):
    for key in (keys or MANIFEST):
        if key not in MANIFEST:
            sys.exit(f'unknown document "{key}" — choose from: {", ".join(MANIFEST)}')
        spec = MANIFEST[key]
        out = convert(load_markdown(spec), ROOT / spec['out'],
                      spec['subtitle'], spec['tagline'], spec['meta'])
        print(f'{key:>16}  ->  {Path(out).name}')


if __name__ == '__main__':
    build(sys.argv[1:] or None)
